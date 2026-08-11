# -*- coding: utf-8 -*-
"""Word 函数说明表生成器。

将解析结果生成为 Word 格式（.docx）的软件详细设计说明书，
包含模块说明、数据结构设计、以及每个函数的说明表。

依赖: python-docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# 配色
COLOR_TITLE = RGBColor(0x1F, 0x38, 0x64)      # 深蓝
COLOR_H1 = RGBColor(0x1F, 0x38, 0x64)
COLOR_H2 = RGBColor(0x2E, 0x54, 0x96)
COLOR_H3 = RGBColor(0x40, 0x40, 0x40)
COLOR_TABLE_HEADER = "1F3864"                 # 表头底色
COLOR_TABLE_HEADER_TEXT = "FFFFFF"
COLOR_ZEBRA = "F2F2F2"                        # 斑马纹


class DocxGenerator:
    """生成 Word 格式的软件详细设计说明书。"""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    # ---------- 样式 ----------

    def _setup_styles(self):
        doc = self.doc
        # 页面：A4，适中页边距
        for section in doc.sections:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            section.left_margin = Cm(2.2)
            section.right_margin = Cm(2.2)
            section.top_margin = Cm(2.4)
            section.bottom_margin = Cm(2.4)

        # Normal 样式
        normal = doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(10.5)
        normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
        pf = normal.paragraph_format
        pf.space_after = Pt(4)
        pf.line_spacing = 1.25

        # 标题样式
        self._setup_heading("Heading 1", 16, COLOR_H1, space_before=16, space_after=8)
        self._setup_heading("Heading 2", 13, COLOR_H2, space_before=12, space_after=6)
        self._setup_heading("Heading 3", 11.5, COLOR_H3, space_before=8, space_after=4)

    def _setup_heading(self, name, size, color, space_before, space_after):
        st = self.doc.styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
        pf = st.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        pf.keep_with_next = True

    # ---------- 基础写入 ----------

    def _para(self, text="", bold=False, size=None, align=None, style=None,
              font=None, color=None, space_after=None):
        p = self.doc.add_paragraph(style=style)
        if align is not None:
            p.alignment = align
        if text:
            run = p.add_run(text)
            run.bold = bold
            if size:
                run.font.size = Pt(size)
            if font:
                run.font.name = font
                run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font)
            if color:
                run.font.color.rgb = color
        if space_after is not None:
            p.paragraph_format.space_after = Pt(space_after)
        return p

    def _cell_para(self, cell, text, bold=False, size=9.5, font="宋体", color=None,
                   align=None):
        """在指定单元格中添加段落（清空默认首段或追加）。"""
        # 若单元格已有空段落则复用，否则新建
        p = cell.paragraphs[0] if cell.paragraphs and not cell.paragraphs[0].runs else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        if align is not None:
            p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = font
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font)
        if color:
            run.font.color.rgb = color
        return p

    def _set_cell_bg(self, cell, hex_color):
        """设置单元格底色。"""
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _set_cell_width(self, cell, cm):
        """设置单元格宽度。"""
        cell.width = Cm(cm)

    def _make_table(self, rows, cols, widths=None):
        table = self.doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.style = "Table Grid"
        if widths:
            for i, w in enumerate(widths):
                for r in table.rows:
                    self._set_cell_width(r.cells[i], w)
        return table

    def _header_row(self, table, headers, widths=None):
        """写表头行。"""
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self._set_cell_bg(cell, COLOR_TABLE_HEADER)
            self._cell_para(cell, h, bold=True, size=9.5, font="黑体",
                            color=RGBColor.from_string(COLOR_TABLE_HEADER_TEXT),
                            align=WD_ALIGN_PARAGRAPH.CENTER)
            if widths:
                self._set_cell_width(cell, widths[j])

    def _zebra(self, table, start=1):
        """斑马纹：给奇数数据行加浅灰底。"""
        for i, row in enumerate(table.rows):
            if i >= start and (i - start) % 2 == 1:
                for cell in row.cells:
                    self._set_cell_bg(cell, COLOR_ZEBRA)

    def _code_block(self, code):
        """等宽字体代码块（带浅灰底）。"""
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(4)
        pf.space_after = Pt(4)
        pf.line_spacing = 1.1
        # 段落底纹
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F7F7F7")
        pPr.append(shd)
        run = p.add_run(code)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
        return p

    # ---------- 文档生成 ----------

    def generate(self, cfile, filename):
        doc = self.doc
        # 封面标题
        self._para("软件详细设计说明书", bold=True, size=20,
                   font="黑体", color=COLOR_TITLE,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
        self._para("C代码自动解析与软件详细设计生成工具",
                   size=12, font="黑体", color=COLOR_H3,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        self._para("（函数说明表与流程分析）",
                   size=10.5, font="宋体", color=COLOR_H3,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

        # 模块说明
        doc.add_heading("1 模块说明", level=1)
        self._module_table(cfile)

        # 数据结构设计
        if cfile.structs or cfile.enums:
            doc.add_heading("2 数据结构设计", level=1)
            sub_no = 1
            for s in cfile.structs:
                self._struct_section(s, sub_no)
                sub_no += 1
            for e in cfile.enums:
                self._enum_section(e, sub_no)
                sub_no += 1

        # 函数设计
        doc.add_heading("3 函数设计", level=1)
        for i, func in enumerate(cfile.functions, 1):
            self._function_section(func, i)
            # 每两个函数之间加分页（保持整洁）
            if i < len(cfile.functions):
                doc.add_page_break()

        doc.save(filename)

    def _module_table(self, cfile):
        rows = [
            ("解析文件", cfile.filename or "未知"),
            ("函数数量", str(len(cfile.functions))),
            ("结构体数量", str(len(cfile.structs))),
            ("枚举数量", str(len(cfile.enums))),
            ("宏定义数量", str(len(cfile.macros))),
            ("包含头文件", ", ".join(cfile.includes) if cfile.includes else "无"),
        ]
        table = self._make_table(len(rows) + 1, 2, widths=[4.5, 11.5])
        self._header_row(table, ["项目", "内容"])
        for i, (k, v) in enumerate(rows, 1):
            c0, c1 = table.rows[i].cells
            c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self._cell_para(c0, k, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self._cell_para(c1, v)
        self._zebra(table)
        self.doc.add_paragraph()

    def _struct_section(self, s, sub_no):
        self.doc.add_heading("2.%d struct %s" % (sub_no, s.name), level=2)
        if not s.members:
            self._para("（无成员）")
            return
        table = self._make_table(len(s.members) + 1, 2, widths=[4.5, 11.5])
        self._header_row(table, ["成员", "类型"])
        for i, m in enumerate(s.members, 1):
            c0, c1 = table.rows[i].cells
            c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self._cell_para(c0, m["name"], align=WD_ALIGN_PARAGRAPH.CENTER)
            c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self._cell_para(c1, m["type"])
        self._zebra(table)
        self.doc.add_paragraph()

    def _enum_section(self, e, sub_no):
        self.doc.add_heading("2.%d enum %s" % (sub_no, e.name), level=2)
        if not e.items:
            self._para("（无枚举项）")
            return
        table = self._make_table(len(e.items) + 1, 2, widths=[4.5, 11.5])
        self._header_row(table, ["枚举项", "值"])
        for i, it in enumerate(e.items, 1):
            c0, c1 = table.rows[i].cells
            c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self._cell_para(c0, it["name"], align=WD_ALIGN_PARAGRAPH.CENTER)
            c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self._cell_para(c1, it["value"])
        self._zebra(table)
        self.doc.add_paragraph()

    def _function_section(self, func, index):
        doc = self.doc
        doc.add_heading("%d.%d 函数 %s" % (3, index, func.name), level=2)

        # 原型
        sig = "%s %s(%s)" % (
            func.return_type,
            func.name,
            ", ".join((p.type + " " + p.name).strip() for p in func.parameters)
            if func.parameters else "void",
        )

        # 解析 Doxygen 字段
        doc_fields = self._parse_doc_fields(func.pre_comment)

        # 概述：@brief 内容（去掉冒号标签），无则"无"
        overview = self._clean_doc_text(doc_fields.get("brief", ""))
        if not overview:
            overview = "无"

        # 参数：每个参数一行
        # 格式：参数名，描述（描述来自 @param，找不到则只写参数名）
        # 无具名参数（含 void）时填"无"
        named_params = [p for p in func.parameters if p.name]
        param_lines = []
        for p in named_params:
            desc = doc_fields.get("param", {}).get(p.name, "")
            if desc:
                param_lines.append("%s，%s" % (p.name, desc))
            else:
                param_lines.append(p.name)
        params_text = "\n".join(param_lines) if param_lines else "无"

        # 返回值
        # 有具体类型时：返回值名/描述（来自 @return），void 则"无"
        ret_raw = doc_fields.get("return", "")
        if func.return_type and func.return_type != "void" and "void" not in func.return_type:
            ret_desc = self._clean_doc_text(ret_raw)
            # 剥掉 "{UINT16}" 类型前缀
            ret_desc = re.sub(r'^\{[^}]*\}\s*', '', ret_desc).strip()
            ret_text = ret_desc if ret_desc else func.return_type
        else:
            ret_text = "无"

        # 引用：调用函数
        ref_text = ", ".join(func.calls) if func.calls else "无"

        # 全局变量引用 / 修改
        g_ref = ", ".join(func.global_refs) if func.global_refs else "无"
        g_mod = ", ".join(func.global_mods) if func.global_mods else "无"

        rows = [
            ("原  型", sig),
            ("概  述", overview),
            ("参  数", params_text),
            ("返回值", ret_text),
            ("引  用", ref_text),
            ("全局变量引用", g_ref),
            ("全局变量修改", g_mod),
        ]
        table = self._make_table(len(rows) + 1, 2, widths=[4.0, 12.0])
        self._header_row(table, ["项目", "内容"])
        for i, (k, v) in enumerate(rows, 1):
            c0, c1 = table.rows[i].cells
            c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self._cell_para(c0, k, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 多行内容（参数）逐行写入
            lines = v.split("\n")
            for j, line in enumerate(lines):
                self._cell_para(c1, line, size=9.5)
        self._zebra(table)
        doc.add_paragraph()

    def _section_title(self, text):
        self._para(text, bold=True, size=10.5, font="黑体", color=COLOR_H3,
                   space_after=2)

    @staticmethod
    def _parse_doc_fields(pre_comment):
        """解析 Doxygen 注释字段。

        支持两种风格：
          - "@brief     描述"（空格分隔）
          - "@brief:    描述"（冒号分隔）
          - "@param {UINT8*} sendBuf，发送缓冲区"（花括号类型 + 名字 + 描述）
          - "@return {UINT16} sendLen，汇总长度"

        返回 dict:
          {"brief": str, "note": str,
           "param": {参数名: 描述}, "return": str}
        """
        result = {"brief": "", "note": "", "param": {}, "return": ""}
        if not pre_comment:
            return result

        # 去掉每行前导星号和行首空白
        lines = []
        for ln in pre_comment.splitlines():
            ln = ln.strip()
            ln = re.sub(r'^/\*+', '', ln)
            ln = re.sub(r'^\*+', '', ln)
            lines.append(ln.strip())

        current_field = None
        for ln in lines:
            m = re.match(r'@(\w+)\s*:?\s*(.*)$', ln)
            if m:
                field = m.group(1)
                content = m.group(2).strip()
                if field in ("brief", "note"):
                    result[field] = content
                    current_field = None
                elif field == "return":
                    result["return"] = content
                    current_field = None
                elif field == "param":
                    current_field = "param"
                    name, desc = DocxGenerator._parse_param_line(content)
                    if name:
                        result["param"][name] = desc
                    # 无名字的参数（如 @param 无）忽略
                else:
                    current_field = None
            elif current_field == "param" and ln:
                # @param 跨行续行（较少见）
                name, desc = DocxGenerator._parse_param_line(ln)
                if name and name not in result["param"]:
                    result["param"][name] = desc

        return result

    @staticmethod
    def _parse_param_line(content):
        """解析 '@param {UINT8*} sendBuf，描述' 行。

        返回 (名字, 描述)。支持：
          - 带花括号类型：'{UINT8*} sendBuf，描述' -> ('sendBuf', '描述')
          - 不带类型：'sendBuf，描述' -> ('sendBuf', '描述')
          - 只有名字：'sendBuf' -> ('sendBuf', '')
        """
        if not content:
            return "", ""
        # 去掉花括号类型 {xxx}
        body = re.sub(r'\{[^}]*\}\s*', '', content).strip()
        if not body:
            return "", ""
        # 逗号分隔：名字，描述
        if "，" in body:
            name, _, desc = body.partition("，")
            return name.strip(), desc.strip()
        if "," in body:
            name, _, desc = body.partition(",")
            return name.strip(), desc.strip()
        return body.strip(), ""

    @staticmethod
    def _clean_doc_text(text):
        """清理文档文本：去掉残留星号、压缩空白。"""
        if not text:
            return ""
        text = re.sub(r'^\s*\*+\s*', '', text)
        text = re.sub(r'\s+', ' ', text).strip()
        return text
